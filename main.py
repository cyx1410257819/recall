import gradio as gr
import torch
import os
import json
import tempfile
import uuid
from pathlib import Path
from flashrag.config import Config
from flashrag.pipeline import ReSearchPipeline
from flashrag.retriever.index_builder import Index_Builder
from flashrag.retriever.utils import load_corpus
import requests
import time
from contextlib import contextmanager

# 创建上传文件和生成索引的目录
UPLOAD_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "uploads")
INDEX_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "indices")
os.makedirs(UPLOAD_DIR, exist_ok=True)
os.makedirs(INDEX_DIR, exist_ok=True)

# 初始化模型和配置的默认值
config_dict = {
    "retrieval_model_path": "/root/cyx/model_weights/Qwen3-Embedding-0.6B",
    "jsonl_path": "/root/cyx/projects/FlashRAG/index/corpus.jsonl",
    "index_path": "/root/cyx/projects/FlashRAG/index/qwen3_emb_Flat.index",
    "generator_model": "/root/cyx/model_weights/ReSearch-Qwen-32B-Instruct",
    "retrieval_topk": 4,
    "dataset_path": "/root/cyx/projects/FlashRAG/datasets",
    "dataset_name": "musique",
    "gpu_id": "4,5,6,7",
    "retrieval_use_fp16": True,
    "split": ["train", "dev"],
    "tensor_parallel_size": 4,
    "retrieval_method": "qwen3",
    "retrieval_query_max_length": 256,
    "retrieval_batch_size": 16,
    "retrieval_pooling_method": "mean",
    "faiss_gpu": False,
    "use_sentence_transformer": False,
    "generator_max_input_len": 32768,
}
# 初始化管道（上传文档时将更新）
pipe = None

# 跟踪当前索引，需要时释放它
current_index = None


def release_gpu_index():
    """释放Faiss索引使用的GPU内存"""
    global pipe, current_index

    try:
        if pipe is not None:
            # 访问包含索引的检索器组件
            if hasattr(pipe, 'retriever') and pipe.retriever is not None:
                # 检查检索器是否有index属性
                if hasattr(pipe.retriever, 'index') and pipe.retriever.index is not None:
                    # 如果索引是GPU索引，先将其移至CPU
                    if hasattr(pipe.retriever.index, 'to_cpu'):
                        print("将Faiss索引从GPU移至CPU...")
                        pipe.retriever.index.to_cpu()

                    # 清除索引引用
                    pipe.retriever.index = None
                    print("从内存中释放Faiss索引")

                # 如果可用，清除其他GPU资源
                if hasattr(pipe.retriever, 'model') and pipe.retriever.model is not None:
                    if hasattr(pipe.retriever.model, 'to_cpu'):
                        pipe.retriever.model.to_cpu()
                    print("将检索器模型移至CPU")

            # 清除管道引用
            pipe = None
            print("清除管道引用")

        # 强制垃圾收集以释放内存
        import gc
        gc.collect()
        if torch.cuda.is_available():
            torch.cuda.empty_cache()
            print("清除CUDA缓存")

        return True
    except Exception as e:
        print(f"释放GPU索引时出错: {str(e)}")
        import traceback
        traceback.print_exc()
        return False


def cleanup_old_indices(keep_recent=5):
    """清理旧索引以节省磁盘空间"""
    try:
        # 获取所有索引目录
        index_dirs = [os.path.join(INDEX_DIR, d) for d in os.listdir(INDEX_DIR)
                      if os.path.isdir(os.path.join(INDEX_DIR, d)) and d.startswith("index_")]

        # 按创建时间排序（最新的优先）
        index_dirs.sort(key=lambda x: os.path.getctime(x), reverse=True)

        # 保留最近的索引
        indices_to_remove = index_dirs[keep_recent:]

        # 删除旧索引
        for index_dir in indices_to_remove:
            try:
                import shutil
                shutil.rmtree(index_dir)
                print(f"已删除旧索引: {index_dir}")
            except Exception as e:
                print(f"删除索引 {index_dir} 时出错: {str(e)}")

        return True
    except Exception as e:
        print(f"清理旧索引时出错: {str(e)}")
        import traceback
        traceback.print_exc()
        return False


def extract_boxed_content(text):
    """从\boxed{}格式中提取内容"""
    import re
    pattern = r'\\boxed{(.*?)}'
    match = re.search(pattern, text)
    return match.group(1) if match else text


def convert_to_jsonl(file_obj, file_name):
    """将上传的文件转换为JSONL格式"""
    import docx
    import PyPDF2
    import pandas as pd
    import io

    print(f"开始转换文件: {file_name}")
    print(f"文件对象类型: {type(file_obj)}")

    # 确保我们有一个有效的文件扩展名
    file_ext = os.path.splitext(file_name)[1].lower() if file_name else ".txt"
    if not file_ext:
        file_ext = ".txt"  # 如果没有扩展名，默认为文本

    temp_jsonl_path = os.path.join(UPLOAD_DIR, f"{uuid.uuid4()}.jsonl")

    print(f"文件扩展名: {file_ext}")
    print(f"临时JSONL路径: {temp_jsonl_path}")

    documents = []

    try:
        if file_ext == '.txt':
            print("处理文本文件...")
            content = file_obj.read() if hasattr(file_obj, 'read') else file_obj
            if isinstance(content, bytes):
                content = content.decode('utf-8')
            print(f"内容长度: {len(content)} 个字符")

            # 按行分割并处理每一行
            lines = [line.strip() for line in content.split('\n') if line.strip()]
            print(f"找到 {len(lines)} 行")

            for i, line in enumerate(lines):
                # 首先尝试解析为JSON
                try:
                    data = json.loads(line)
                    if isinstance(data, dict):
                        # 如果已经是正确的格式，使用它
                        if all(k in data for k in ["id", "title", "contents"]):
                            documents.append(data)
                        else:
                            # 否则，创建一个新文档
                            documents.append({
                                "id": i,
                                "title": data.get("title", f"行 {i + 1}"),
                                "contents": data.get("contents", line)
                            })
                    else:
                        # 如果不是dict，使用该行作为内容
                        documents.append({
                            "id": i,
                            "title": f"行 {i + 1}",
                            "contents": line
                        })
                except json.JSONDecodeError:
                    # 如果不是JSON，使用该行作为内容
                    documents.append({
                        "id": i,
                        "title": f"行 {i + 1}",
                        "contents": line
                    })

        elif file_ext == '.json':
            print("处理JSON文件...")
            if isinstance(file_obj, bytes):
                data = json.loads(file_obj.decode('utf-8'))
            else:
                data = json.load(file_obj)

            if isinstance(data, list):
                print(f"处理包含 {len(data)} 个项目的JSON数组")
                for i, item in enumerate(data):
                    if isinstance(item, dict):
                        # 如果已经是正确的格式，使用它
                        if all(k in item for k in ["id", "title", "contents"]):
                            documents.append(item)
                        else:
                            # 否则，创建一个新文档
                            documents.append({
                                "id": i,
                                "title": item.get("title", f"项目 {i + 1}"),
                                "contents": item.get("contents", json.dumps(item, ensure_ascii=False))
                            })
                    else:
                        # 如果不是dict，转换为字符串
                        documents.append({
                            "id": i,
                            "title": f"项目 {i + 1}",
                            "contents": str(item)
                        })
            else:
                # 单个JSON对象
                if isinstance(data, dict):
                    # 如果已经是正确的格式，使用它
                    if all(k in data for k in ["id", "title", "contents"]):
                        documents.append(data)
                    else:
                        # 否则，创建一个新文档
                        documents.append({
                            "id": 0,
                            "title": data.get("title", "文档"),
                            "contents": data.get("contents", json.dumps(data, ensure_ascii=False))
                        })
                else:
                    # 如果不是dict，转换为字符串
                    documents.append({
                        "id": 0,
                        "title": "文档",
                        "contents": str(data)
                    })

        elif file_ext == '.jsonl':
            print("处理JSONL文件...")
            content = file_obj.read() if hasattr(file_obj, 'read') else file_obj
            if isinstance(content, bytes):
                content = content.decode('utf-8')

            # 处理每一行作为JSON
            for i, line in enumerate(content.split('\n')):
                line = line.strip()
                if not line:
                    continue

                try:
                    data = json.loads(line)
                    if isinstance(data, dict):
                        # 如果已经是正确的格式，使用它
                        if all(k in data for k in ["id", "title", "contents"]):
                            documents.append(data)
                        else:
                            # 否则，创建一个新文档
                            documents.append({
                                "id": i,
                                "title": data.get("title", f"行 {i + 1}"),
                                "contents": data.get("contents", json.dumps(data, ensure_ascii=False))
                            })
                    else:
                        # 如果不是dict，转换为字符串
                        documents.append({
                            "id": i,
                            "title": f"行 {i + 1}",
                            "contents": str(data)
                        })
                except json.JSONDecodeError:
                    # 如果不是有效的JSON，使用该行作为内容
                    documents.append({
                        "id": i,
                        "title": f"行 {i + 1}",
                        "contents": line
                    })

        elif file_ext == '.docx':
            print("处理Word文档...")
            doc_stream = io.BytesIO(file_obj) if isinstance(file_obj, bytes) else file_obj
            doc = docx.Document(doc_stream)
            print(f"找到 {len(doc.paragraphs)} 个段落")

            current_title = None
            current_content = []

            for i, para in enumerate(doc.paragraphs):
                if para.style.name.startswith('Heading'):
                    # 保存前面的部分（如果存在）
                    if current_title and current_content:
                        documents.append({
                            "id": len(documents),
                            "title": current_title,
                            "contents": f"{current_title} {' '.join(current_content)}"
                        })
                    current_title = para.text.strip()
                    current_content = []
                elif para.text.strip():
                    current_content.append(para.text.strip())

            # 添加最后一部分
            if current_title and current_content:
                documents.append({
                    "id": len(documents),
                    "title": current_title,
                    "contents": f"{current_title} {' '.join(current_content)}"
                })

        elif file_ext == '.pdf':
            print("处理PDF文件...")
            pdf_stream = io.BytesIO(file_obj) if isinstance(file_obj, bytes) else file_obj
            pdf_reader = PyPDF2.PdfReader(pdf_stream)
            print(f"找到 {len(pdf_reader.pages)} 页")

            for i, page in enumerate(pdf_reader.pages):
                text = page.extract_text().strip()
                if text:
                    # 将文本拆分为段落
                    paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]
                    for j, para in enumerate(paragraphs):
                        # 尝试提取问题和答案
                        lines = [line.strip() for line in para.split('\n') if line.strip()]
                        if len(lines) >= 2:
                            title = lines[0]
                            contents = ' '.join(lines[1:])
                            documents.append({
                                "id": len(documents),
                                "title": title,
                                "contents": f"{title} {contents}"
                            })

        elif file_ext == '.csv':
            print("处理CSV文件...")
            if isinstance(file_obj, bytes):
                csv_stream = io.StringIO(file_obj.decode('utf-8'))
            else:
                csv_stream = file_obj

            df = pd.read_csv(csv_stream)
            print(f"找到 {len(df)} 行")

            # 尝试识别问题和答案列
            question_col = None
            answer_col = None

            # 查找可能包含问题和答案的列
            for col in df.columns:
                if 'question' in col.lower() or 'q' in col.lower():
                    question_col = col
                elif 'answer' in col.lower() or 'a' in col.lower():
                    answer_col = col

            # 如果找到了问题和答案列，使用它们
            if question_col and answer_col:
                for i, row in df.iterrows():
                    question = str(row[question_col]).strip()
                    answer = str(row[answer_col]).strip()
                    if question and answer:
                        documents.append({
                            "id": i,
                            "title": question,
                            "contents": f"{question} {answer}"
                        })
            else:
                # 否则，将每一行视为问题-答案对
                for i, row in df.iterrows():
                    row_data = [f"{col}: {val}" for col, val in row.items() if pd.notna(val)]
                    if len(row_data) >= 2:
                        title = row_data[0]
                        contents = ' '.join(row_data[1:])
                        documents.append({
                            "id": i,
                            "title": title,
                            "contents": f"{title} {contents}"
                        })

        else:
            raise ValueError(f"不支持的文件扩展名: {file_ext}")

        # 确保我们至少有一个文档
        if not documents:
            raise ValueError(f"在 {file_ext} 文件中未找到有效内容")

        print(f"提取了 {len(documents)} 个文档")

        # 写入JSONL文件
        valid_docs = []
        with open(temp_jsonl_path, 'w', encoding='utf-8') as f:
            for doc in documents:
                # 检查是否存在所有必填字段且有效
                if ("id" in doc and
                        "title" in doc and doc["title"].strip() and
                        "contents" in doc and doc["contents"].strip()):

                    # 确保内容包含标题
                    if not doc["contents"].startswith(doc["title"]):
                        doc["contents"] = f"{doc['title']} {doc['contents']}"

                    json.dump(doc, f, ensure_ascii=False)
                    f.write('\n')
                    valid_docs.append(doc)

        if not valid_docs:
            raise ValueError("没有有效的文档可以写入JSONL文件")

        print(f"成功将 {len(valid_docs)} 个文档写入JSONL文件")
        return temp_jsonl_path, len(valid_docs)

    except Exception as e:
        import traceback
        error_msg = f"转换文件时出错: {str(e)}\n{traceback.format_exc()}"
        print(error_msg)
        if os.path.exists(temp_jsonl_path):
            os.remove(temp_jsonl_path)
        return None, 0


def build_index(jsonl_path):
    """从JSONL文件构建索引"""
    try:
        # 验证JSONL文件存在且有内容
        if not os.path.exists(jsonl_path):
            raise FileNotFoundError(f"未找到JSONL文件: {jsonl_path}")

        # 创建唯一的索引名称
        index_name = f"index_{uuid.uuid4()}"
        index_path = os.path.join(INDEX_DIR, index_name)
        os.makedirs(index_path, exist_ok=True)

        print(f"在 {index_path} 构建索引")
        print(f"使用JSONL文件: {jsonl_path}")

        # 使用所有必需的参数创建索引构建器
        index_builder = Index_Builder(
            retrieval_method=config_dict["retrieval_method"],
            model_path=config_dict["retrieval_model_path"],
            corpus_path=jsonl_path,
            save_dir=index_path,
            max_length=config_dict["retrieval_query_max_length"],
            batch_size=config_dict["retrieval_batch_size"],
            use_fp16=config_dict["retrieval_use_fp16"],
            pooling_method=config_dict["retrieval_pooling_method"],
            instruction=None,
            faiss_type="Flat",
            embedding_path=None,
            save_embedding=False,
            faiss_gpu=config_dict["faiss_gpu"],
            use_sentence_transformer=config_dict["use_sentence_transformer"],
            bm25_backend="bm25s",
            index_modal="all"
        )

        # 构建索引
        print(f"开始索引构建过程...")
        index_builder.build_index()
        print(f"索引构建成功完成")

        # 获取实际的索引文件路径
        index_file = os.path.join(index_path, f"{config_dict['retrieval_method']}_Flat.index")
        if not os.path.exists(index_file):
            raise FileNotFoundError(f"未创建索引文件: {index_file}")

        print(f"索引文件创建于: {index_file}")
        return index_file, True
    except Exception as e:
        import traceback
        error_msg = f"构建索引时出错: {str(e)}\n{traceback.format_exc()}"
        print(error_msg)
        return None, False


def process_uploaded_file(single_file, multi_files, dir_files):
    """通过pipe的API处理上传文件并构建索引"""
    if not any([single_file, multi_files, dir_files]):
        return "未上传文件", None, None

    print("通过pipe API处理上传的文件...")

    release_gpu_index()

    # 收集所有文件路径
    file_paths = []
    if single_file:
        file_paths.append(single_file)
    if multi_files:
        file_paths.extend(multi_files)
    if dir_files:
        file_paths.extend(dir_files)

    # 只支持单文件上传API演示，如需多文件可循环调用
    standard_data = []
    upload_files = []
    for file_path in file_paths:
        upload_files.append(
            ("files", (os.path.basename(file_path), open(file_path, "rb"), "application/octet-stream"))
        )
    try:
        resp = requests.post("http://localhost:8001/process", files=upload_files, timeout=60000)
        resp.raise_for_status()
        result = resp.json()
        standard_data.extend(result.get("data", []))
    except Exception as e:
        print(f"Pipe API错误: {e}")
        return f"Pipe API错误: {e}", None, None

    if not standard_data:
        return "未提取到有效内容", None, None

    # 写入JSONL
    combined_jsonl_path = os.path.join(UPLOAD_DIR, f"combined_{uuid.uuid4()}.jsonl")
    with open(combined_jsonl_path, 'w', encoding='utf-8') as f:
        for doc in standard_data:
            json.dump(doc, f, ensure_ascii=False)
            f.write('\n')

    # 构建索引
    index_path, success = build_index(combined_jsonl_path)
    if not success:
        return f"索引构建失败", None, None

    # 更新pipeline
    global pipe, config_dict, current_index
    config_dict["index_path"] = index_path
    config_dict["corpus_path"] = combined_jsonl_path
    config = Config('/root/cyx/projects/ReSearch/basic_config.yaml', config_dict=config_dict)
    json_path = '/root/cyx/projects/ReSearch/data.json'
    doc_path = combined_jsonl_path
    pipe = ReSearchPipeline(config, apply_chat=True, json_path=json_path, doc_path=doc_path)
    current_index = index_path

    cleanup_old_indices()
    return f"成功处理{len(file_paths)}个文件", index_path, combined_jsonl_path


def process_query(query: str, show_reasoning: bool = True, use_background: bool = True):
    """
    处理单个查询并返回完整回答
    """
    global pipe

    print(f"\n=== 开始process_query ===")
    print(f"查询: {query}")
    print(f"显示推理: {show_reasoning}")
    print(f"使用背景: {use_background}")

    if pipe is None:
        # 如果尚未初始化，使用默认文档初始化
        try:
            default_jsonl_path = os.path.join(UPLOAD_DIR, "default.jsonl")
            print(f"使用默认文档初始化: {default_jsonl_path}")

            # 创建默认索引（如果不存在）
            print(f"从 {default_jsonl_path} 创建默认索引")
            index_path, success = build_index(default_jsonl_path)
            if not success:
                print("创建默认索引失败")
                return "错误: 初始化默认文档失败。请先上传文档。"
            print(f"在 {index_path} 成功创建默认索引")

            # 使用默认文档初始化管道
            config_dict["index_path"] = index_path
            config_dict["corpus_path"] = default_jsonl_path
            config = Config('/root/cyx/projects/ReSearch/basic_config.yaml',
                            config_dict=config_dict)
            json_path = '/root/cyx/projects/ReSearch/data.json'
            doc_path = default_jsonl_path
            pipe = ReSearchPipeline(config, apply_chat=True, use_background=use_background, json_path=json_path,
                                    doc_path=doc_path)
            print("成功使用默认文档初始化管道")

        except Exception as e:
            print(f"初始化默认文档时出错: {str(e)}")
            import traceback
            traceback.print_exc()
            return f"错误: 初始化默认文档失败: {str(e)}。请先上传文档。"

    try:
        print(f"处理查询: {query}")

        class SimpleDataset:
            def __init__(self, question):
                self.question = [question]
                self.output = {}
                self._idx = 0
                self.pred = []
                self.choices = []
                self.prompt = []

            def update_output(self, key, value):
                self.output[key] = value
                print(f"更新输出 {key}: {value}")

            def __iter__(self):
                self._idx = 0
                return self

            def __next__(self):
                if self._idx < len(self.question):
                    item = type('DatasetItem', (), {
                        'question': self.question[self._idx],
                        'update_output': lambda key, value: self.update_output(key, value)
                    })
                    self._idx += 1
                    return item
                raise StopIteration

            def __len__(self):
                return len(self.question)

            def save(self, save_path):
                pass

            def to_dict(self):
                return {
                    'question': self.question,
                    'output': self.question,
                    'pred': self.pred,
                    'choices': self.choices,
                    'prompt': self.prompt
                }

        test_data = SimpleDataset(query)
        pipe.use_background = use_background

        print("开始生成...")

        # 处理生成器返回的内容
        try:
            for item in test_data:
                # 使用stream_generate获取流式结果（如果可用）
                if hasattr(pipe, 'stream_generate'):
                    response = ""
                    for chunk in pipe.stream_generate(item.question, item.question, max_length=4096):
                        response += chunk
                        # 提取<|im_start|>assistant后面的内容
                        if "<|im_start|>assistant" in response:
                            response = response.split("<|im_start|>assistant")[1].strip()
                        # 每生成一段就更新一次
                        yield response
                        time.sleep(0.05)  # 控制更新频率
                    return
                else:
                    # 如果没有stream_generate方法，使用run_item
                    response = pipe.run_item(item)
                    if response:
                        # 提取<|im_start|>assistant后面的内容
                        if "<|im_start|>assistant" in response:
                            response = response.split("<|im_start|>assistant")[1].strip()
                        print(f"生成的回答: {response}")
                        yield response
                        return
                    else:
                        print("run_item没有生成回答")
                        yield "抱歉，我无法生成回答。请尝试重新表述您的问题。"
                        return

        except Exception as e:
            print(f"生成时出错: {str(e)}")
            import traceback
            traceback.print_exc()
            yield f"错误: 生成过程中出错: {str(e)}"

    except Exception as e:
        print(f"process_query中出错: {str(e)}")
        import traceback
        traceback.print_exc()
        yield f"错误: 处理查询时出错: {str(e)}"


# 创建Gradio界面
with gr.Blocks(title="文档搜索问答系统 (流式输出)") as demo:
    gr.Markdown("""
    # 文档搜索问答系统 (流式输出)
    上传您的文档并对其进行提问。该系统将使用先进的检索增强生成来逐步回答您的问题。
    """)

    with gr.Row():
        with gr.Column():
            # 文件上传组件
            file_input_single = gr.File(
                label="上传单文件",
                file_types=[".txt", ".docx", ".pdf", ".csv", ".json"],
                type="filepath",  # 改为filepath以获取实际文件路径
                file_count="single",
            )
            file_input_multi = gr.File(
                label="上传多文件",
                file_types=[".txt", ".docx", ".pdf", ".csv", ".json"],
                type="filepath",  # 改为filepath以获取实际文件路径
                file_count="multiple",
            )
            file_input_dir = gr.File(
                label="上传文件夹",
                file_types=[".txt", ".docx", ".pdf", ".csv", ".json"],
                type="filepath",  # 改为filepath以获取实际文件路径
                file_count="directory",
            )

            file_names = gr.Textbox(label="文件名", visible=False)
            upload_button = gr.Button("文档处理")
            upload_output = gr.Textbox(label="上传状态", interactive=False)

            # 问题输入
            question_input = gr.Textbox(
                label="你的问题",
                placeholder="在这里输入你的问题…",
                lines=3
            )
            show_reasoning = gr.Checkbox(
                label="展示推理过程",
                value=True
            )
            use_background = gr.Checkbox(
                label="使用背景调查",
                value=True
            )
            submit_btn = gr.Button("开始回答")

        with gr.Column():
            answer_output = gr.Textbox(
                label="回答",
                lines=10,
                interactive=False,
                value="等待你的问题……"  # 添加默认值
            )


    # 处理文件选择
    def on_file_select(single_file, multi_files, dir_files):
        """处理文件选择"""
        if not any([single_file, multi_files, dir_files]):
            return "未选择文件", ""

        file_names = []
        if single_file:
            file_names.append(os.path.basename(single_file))
        if multi_files:
            file_names.extend([os.path.basename(f) for f in multi_files])
        if dir_files:
            file_names.extend([os.path.basename(f) for f in dir_files])

        return ", ".join(file_names), ", ".join([os.path.splitext(f)[1] for f in file_names])


    # 更新文件输入更改处理程序
    file_input_single.change(
        fn=lambda x: on_file_select(x, None, None),
        inputs=[file_input_single],
        outputs=[file_names, upload_output]
    )

    file_input_multi.change(
        fn=lambda x: on_file_select(None, x, None),
        inputs=[file_input_multi],
        outputs=[file_names, upload_output]
    )

    file_input_dir.change(
        fn=lambda x: on_file_select(None, None, x),
        inputs=[file_input_dir],
        outputs=[file_names, upload_output]
    )

    # 处理上传按钮点击
    upload_button.click(
        fn=process_uploaded_file,
        inputs=[file_input_single, file_input_multi, file_input_dir],
        outputs=[upload_output, gr.State(), gr.State()]
    )


    # 处理提交按钮点击 - 使用自定义流式实现
    def handle_submit(question, show_reasoning, use_background):
        if not question.strip():
            yield "请输入问题。"
            return

        # 初始化回答
        answer = ""

        # 使用生成器逐步获取回答
        try:
            for chunk in process_query(question, show_reasoning, use_background):
                answer = chunk
                yield answer
                time.sleep(0.01)  # 控制更新频率
        except Exception as e:
            yield f"错误: {str(e)}"


    submit_btn.click(
        fn=handle_submit,
        inputs=[question_input, show_reasoning, use_background],
        outputs=answer_output
    )

    gr.Examples(
        examples=[
            ["赛科希德的竞争对手有那些？"],
            ["北京谊安的总部地址"],
            ["湖南现任的省委常委都有那些人？"],
            ["大众2024年的汽车销量是多少？"],
            ["北京最大的商城在哪里？"],
            ["公司的核心价值观是什么？"],
            ["赛科希德的董事长是谁？"]
        ],
        inputs=question_input
    )

# 启动应用程序
if __name__ == "__main__":
    # 注册应用关闭时运行的函数
    import atexit

    atexit.register(release_gpu_index)

    # 启动应用
    demo.launch(server_name="0.0.0.0", server_port=8004, share=True)


